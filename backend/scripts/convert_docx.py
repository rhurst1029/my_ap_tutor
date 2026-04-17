import json, sys, os, re
from pathlib import Path
from docx import Document

DOCX_PATH   = Path(__file__).parent.parent / "bella_ap_csa_test.docx"
GUIDE_PATH  = Path(__file__).parent.parent / "bella_tutor_guide.docx"
OUTPUT_PATH = Path(__file__).parent.parent / "data/tests/ap_csa_test_1.json"

VALID_TOPICS = [
    "variables_and_types", "operators", "conditionals", "loops",
    "arrays", "2d_arrays", "methods", "parameter_passing", "strings",
    "classes_and_objects", "inheritance", "polymorphism", "recursion",
    "arraylist", "searching_sorting", "interfaces"
]

# Heuristic topic mapping from unit/concept in answer key
CONCEPT_TOPIC_MAP = {
    "String": ["strings"],
    "substring": ["strings"],
    "Integer division": ["operators", "variables_and_types"],
    "type widening": ["variables_and_types"],
    "Boolean logic": ["conditionals", "operators"],
    "Loop tracing": ["loops"],
    "modulus": ["loops", "operators"],
    "Nested loop": ["loops"],
    "Object instantiation": ["classes_and_objects", "methods"],
    "method call": ["methods", "classes_and_objects"],
    "Encapsulation": ["classes_and_objects"],
    "data hiding": ["classes_and_objects"],
    "Array traversal": ["arrays"],
    "max": ["arrays"],
    "ArrayList": ["arraylist"],
    "2D array": ["2d_arrays"],
    "indexing": ["2d_arrays"],
}


def inspect_document(path: Path):
    doc = Document(str(path))
    print(f"\n=== PARAGRAPHS in {path.name} ===")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            fonts = list({run.font.name for run in para.runs if run.font.name})
            style_name = para.style.name if para.style else "None"
            print(f"[P{i:03d}] style={style_name!r:30s} fonts={fonts} text={para.text[:120]!r}")
    print(f"\n=== TABLES in {path.name} ===")
    for i, table in enumerate(doc.tables):
        print(f"[Table {i}] {len(table.rows)} rows x {len(table.columns)} cols")
        for row in table.rows:
            print("  row:", [c.text.strip()[:40] for c in row.cells])


def get_paragraph_fonts(para) -> list:
    return list({run.font.name for run in para.runs if run.font.name})


def is_code_paragraph(para) -> bool:
    fonts = get_paragraph_fonts(para)
    return "Courier New" in fonts


def parse_answer_key(guide_doc: Document) -> dict:
    """Extract CSA answer key from the FIRST matching table in tutor guide (Table 1)."""
    answer_key = {}
    for table in guide_doc.tables:
        if len(table.rows) >= 2 and len(table.columns) >= 2:
            header = [c.text.strip() for c in table.rows[0].cells]
            if "#" in header and "Answer" in header:
                num_col = header.index("#")
                ans_col = header.index("Answer")
                concept_col = header.index("Concept Tested") if "Concept Tested" in header else None
                unit_col = header.index("Unit") if "Unit" in header else None
                for row in table.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    if len(cells) > max(num_col, ans_col):
                        q_num = cells[num_col].strip()
                        ans_text = cells[ans_col].strip()
                        concept = cells[concept_col].strip() if concept_col is not None and len(cells) > concept_col else ""
                        unit = cells[unit_col].strip() if unit_col is not None and len(cells) > unit_col else ""
                        # Extract letter: 'A — "put"' -> 'A'
                        letter_match = re.match(r'^([A-D])', ans_text)
                        if letter_match:
                            letter = letter_match.group(1)
                            answer_key[q_num] = {
                                "answer": letter,
                                "concept": concept,
                                "unit": unit,
                                "full_answer_text": ans_text,
                            }
                # Stop after FIRST matching table (CSA, not Calc BC)
                break
    return answer_key


def parse_guiding_questions(guide_doc: Document) -> dict:
    """
    Extract guiding questions per CSA question number from tutor guide.
    Returns dict: { "1": [ "Ask: ...", ... ], ... }
    Looks only in the PART 2 (CSA) section, stops at PART 3 (Calc BC).
    """
    guiding = {}
    current_q = None
    q_counter = 0
    in_csa_section = False
    in_probing = False

    paragraphs = list(guide_doc.paragraphs)
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""

        # Enter CSA section at PART 2
        if "PART 2" in text and "Computer Science" in text:
            in_csa_section = True
            continue

        # Stop at PART 3
        if "PART 3" in text:
            break

        if not in_csa_section:
            continue

        # Each "What Each Wrong Answer Reveals:" starts a new question block
        if text.startswith("What Each Wrong Answer Reveals:"):
            q_counter += 1
            current_q = str(q_counter)
            in_probing = False

        elif text.startswith("Probing Questions"):
            in_probing = True

        elif in_probing and style_name == "List Paragraph":
            if current_q is not None:
                if current_q not in guiding:
                    guiding[current_q] = []
                guiding[current_q].append(text)

        elif text.startswith("Knowledge Signal:"):
            in_probing = False

    return guiding


def infer_topics_from_concept(concept: str) -> list:
    """Map concept string to topic tags."""
    for key, topics in CONCEPT_TOPIC_MAP.items():
        if key.lower() in concept.lower():
            return topics
    return ["variables_and_types"]


def extract_questions(test_doc: Document, answer_key: dict, guiding_questions: dict) -> list:
    """Parse test document paragraphs into question objects."""
    questions = []
    current_q = None
    code_lines = []
    collecting_code = False

    paragraphs = list(test_doc.paragraphs)

    for i, para in enumerate(paragraphs):
        text = para.text.strip()

        # Check for question header: "Question N   [Unit X — Topic]"
        q_match = re.match(r'^Question\s+(\d+)\s*\[', text)
        if q_match:
            # Finalize the previous question
            if current_q is not None:
                if code_lines:
                    current_q["code_block"] = "\n".join(code_lines)
                questions.append(current_q)

            q_num = q_match.group(1)
            bracket_match = re.search(r'\[(.+?)\]', text)
            topic_hint = bracket_match.group(1) if bracket_match else ""

            # Get answer info from key
            ans_info = answer_key.get(q_num, {})
            concept = ans_info.get("concept", topic_hint)

            current_q = {
                "id": f"q{q_num}",
                "type": None,
                "topic_tags": infer_topics_from_concept(concept or topic_hint),
                "difficulty": "medium",
                "text": None,
                "code_block": None,
                "options": [],
                "answer_key": ans_info.get("answer", ""),
                "explanation": f"Correct answer: {ans_info.get('full_answer_text', '')}. Concept: {concept}",
                "guiding_questions": [],
                "execution_trace": None,
                "_q_num": q_num,
            }
            code_lines = []
            collecting_code = False
            continue

        if current_q is None:
            continue

        # Code paragraph (Courier New font)
        if is_code_paragraph(para) and text:
            code_lines.append(text)
            collecting_code = True
            continue

        # Non-code paragraph after code block — finalize code block
        if collecting_code and not is_code_paragraph(para):
            if text:  # non-empty non-code: ends code block
                collecting_code = False
            # If empty: don't end yet (allow blank lines inside code blocks)

        # MC option: "(A)  put"
        mc_match = re.match(r'^\(([A-D])\)\s+(.+)', text)
        if mc_match and text:
            letter = mc_match.group(1)
            option_text = mc_match.group(2).strip()
            current_q["options"].append(f"{letter}) {option_text}")
            continue

        # Skip administrative text
        if text.startswith("Work / Reasoning:") or text.startswith("—") or text.startswith("Student:") or text.startswith("Instructions:") or text.startswith("AP Computer Science"):
            continue

        # Question text (question stem)
        if current_q["text"] is None and text and not is_code_paragraph(para):
            current_q["text"] = text

    # Don't forget the last question
    if current_q is not None:
        if code_lines:
            current_q["code_block"] = "\n".join(code_lines)
        questions.append(current_q)

    # Post-process: assign types, guiding questions
    for q in questions:
        q_num_str = q["_q_num"]

        # Build guiding questions list
        gq_texts = guiding_questions.get(q_num_str, [])
        gq_list = []
        for idx, gq_text in enumerate(gq_texts[:3]):  # cap at 3 per question
            gq_list.append({
                "id": f"q{q_num_str}_g{idx+1}",
                "text": gq_text,
                "intent": f"Probe understanding for question {q_num_str}"
            })
        q["guiding_questions"] = gq_list

        # Determine question type
        if q["code_block"] and q["options"]:
            q["type"] = "code_trace"
        elif q["options"]:
            q["type"] = "multiple_choice"
        else:
            q["type"] = "free_response"

        # Difficulty heuristic
        if q["type"] == "code_trace":
            q["difficulty"] = "hard"
        elif q["type"] == "free_response":
            q["difficulty"] = "hard"
        else:
            q["difficulty"] = "medium"

        # Remove internal key
        del q["_q_num"]

    return questions


def build_test_json(questions: list) -> dict:
    return {
        "$schema": "ap_csa_question_schema_v1",
        "test_id": "ap_csa_test_1",
        "title": "AP Computer Science A — Diagnostic Test",
        "created_at": "2026-04-11",
        "source": "converted_docx",
        "questions": questions,
    }


if __name__ == "__main__":
    if "--inspect" in sys.argv:
        inspect_document(DOCX_PATH)
        inspect_document(GUIDE_PATH)
        sys.exit(0)

    print("Loading documents...")
    test_doc  = Document(str(DOCX_PATH))
    guide_doc = Document(str(GUIDE_PATH))

    print("Parsing answer key from tutor guide...")
    answer_key = parse_answer_key(guide_doc)
    print(f"  Found {len(answer_key)} answers:")
    for qn, info in sorted(answer_key.items(), key=lambda x: int(x[0])):
        print(f"    Q{qn}: {info['answer']} — {info['full_answer_text']}")

    print("Parsing guiding questions from tutor guide...")
    guiding_questions = parse_guiding_questions(guide_doc)
    for qn, gqs in sorted(guiding_questions.items(), key=lambda x: int(x[0])):
        print(f"  Q{qn}: {len(gqs)} guiding questions")

    print("Extracting questions from test document...")
    questions = extract_questions(test_doc, answer_key, guiding_questions)
    print(f"  Extracted {len(questions)} questions")

    # Try Claude topic tagging if API key available
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if api_key:
        print("Assigning topic tags via Claude...")
        try:
            import anthropic
            client = anthropic.Anthropic(api_key=api_key)
            for q in questions:
                content = (q["text"] or "") + ("\n" + q["code_block"] if q["code_block"] else "")
                msg = client.messages.create(
                    model="claude-opus-4-6",
                    max_tokens=200,
                    messages=[{"role": "user", "content":
                        f"Given this AP CSA question, return a JSON array of topic tags "
                        f"from this list ONLY: {VALID_TOPICS}\n"
                        f"Question: {content}\n"
                        f"Return only the JSON array, no explanation, no markdown."}]
                )
                raw = msg.content[0].text.strip().strip("```json").strip("```").strip()
                tags = json.loads(raw)
                if tags:
                    q["topic_tags"] = tags
                print(f"  {q['id']}: {tags}")
        except Exception as e:
            print(f"  Warning: Claude tagging failed ({e}), using heuristic tags")
    else:
        print("ANTHROPIC_API_KEY not set — using heuristic topic tags")

    test_data = build_test_json(questions)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PATH.write_text(json.dumps(test_data, indent=2))
    print(f"\nWrote {len(questions)} questions to {OUTPUT_PATH}")

    # Summary
    print("\nQuestion Summary:")
    for q in questions:
        print(f"  {q['id']}: type={q['type']}, answer={q['answer_key']}, "
              f"options={len(q.get('options', []))}, guiding={len(q['guiding_questions'])}, "
              f"code={'yes' if q['code_block'] else 'no'}, topics={q['topic_tags']}")
